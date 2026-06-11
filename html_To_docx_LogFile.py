#!/usr/bin/env python3
"""
Converts HTML files containing historyElement divs into individual .docx files.
 
Usage:
    python html_to_docx.py input.html [output_dir]
    python html_to_docx.py input_folder/ [output_dir]
 
Each historyElement becomes one .docx file named:
    "{rank}. Başlık {first language title}.docx"
 
Each .docx contains, per language:
    - Heading 1: language title
    - Body: content (bold/italic/paragraph structure preserved)
    - Heading 4: "Meta Description:"
    - Body: meta description
    - Heading 4: "Meta Keywords:"
    - Body: keywords
    - 3 empty lines (between languages)
 
Requirements:
    pip install python-docx beautifulsoup4
"""
 
import os
import re
import sys
from pathlib import Path
import html
 
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag

from docx import Document as DocumentFactory
from docx.document import Document as DocxDocument
 
 
# ---------------------------------------------------------------------------
# HTML → docx content helpers
# ---------------------------------------------------------------------------
 
BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
              "ul", "ol", "li", "blockquote", "pre"}
 
 
def _add_inline(paragraph, node, bold=False, italic=False, underline=False):
    """Recursively walk inline nodes and add styled runs to *paragraph*."""
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline
        return
 
    if not isinstance(node, Tag):
        return
 
    tag = node.name.lower() if node.name else ""
 
    if tag == "br":
        paragraph.add_run("\n")
        return
 
    new_bold = bold or tag in ("b", "strong")
    new_italic = italic or tag in ("i", "em")
    new_underline = underline or tag == "u"
 
    for child in node.children:
        _add_inline(paragraph, child, new_bold, new_italic, new_underline)
 
 
def add_html_to_doc(doc, raw_html: str):
    """
    Convert *raw_html* (textarea inner HTML) to docx paragraphs.
 
    Preserves:
      • <b>/<strong> → bold
      • <i>/<em>     → italic
      • <u>          → underline
      • <p>          → separate paragraphs
      • <br>         → line break within a paragraph
      • <ul>/<ol>    → bulleted / numbered list
      • Plain text   → single paragraph (leading/trailing whitespace stripped)
    """
    raw_html = html.unescape(raw_html)
    soup = BeautifulSoup(raw_html, "html.parser")
 
    # Check whether the content contains any block-level elements.
    has_blocks = any(
        isinstance(c, Tag) and c.name and c.name.lower() in BLOCK_TAGS
        for c in soup.children
    )
 
    if not has_blocks:
        # ── Plain / inline content ──────────────────────────────────────────
        # Split on <br> tags to create separate paragraphs.
        raw_str = str(soup)
        parts = re.split(r"<br\s*/?>", raw_str, flags=re.IGNORECASE)
        for part in parts:
            part = part.strip()
            if not part:
                continue
            p = doc.add_paragraph()
            part_soup = BeautifulSoup(part, "html.parser")
            _add_inline(p, part_soup)
        return
 
    # ── Block-level content ─────────────────────────────────────────────────
    for child in soup.children:
        if isinstance(child, NavigableString):
            text = child.strip()
            if text:
                p = doc.add_paragraph()
                p.add_run(text)
            continue
 
        if not isinstance(child, Tag):
            continue
 
        tag = child.name.lower() if child.name else ""
 
        if tag in ("p", "div", "blockquote", "pre"):
            p = doc.add_paragraph()
            _add_inline(p, child)
 
        elif tag in ("ul", "ol"):
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in child.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                _add_inline(p, li)
 
        elif tag == "li":
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, child)
 
        elif re.match(r"h[1-6]", tag):
            level = int(tag[1])
            p = doc.add_heading("", level=level)
            _add_inline(p, child)
 
        else:
            # Fallback: treat as inline wrapped in a paragraph
            p = doc.add_paragraph()
            _add_inline(p, child)
 
 
# ---------------------------------------------------------------------------
# Parsing historyElement
# ---------------------------------------------------------------------------
 
def parse_history_element(element: Tag) -> list[dict]:
    """
    Return a list of language dicts from a historyElement <div>.
 
    Each dict has: title, content_html, metadesc_html, keywords_html
    """
    inner_div = element.find("div")
    if not inner_div:
        return []
 
    languages = []
    children = list(inner_div.children)
    i = 0
 
    while i < len(children):
        child = children[i]
 
        # A new language section starts at an <h3>
        if isinstance(child, Tag) and child.name == "h3":
            lang = {
                "title": child.get_text(strip=True),
                "content_html": "",
                "metadesc_html": "",
                "keywords_html": "",
            }
            i += 1
            textarea_idx = 0  # 0=content, 1=metadesc, 2=keywords
 
            while i < len(children):
                c = children[i]
 
                # Next language starts
                if isinstance(c, Tag) and c.name == "h3":
                    break
 
                if isinstance(c, Tag) and c.name == "textarea":
                    inner = c.decode_contents()   # raw innerHTML
                    if textarea_idx == 0:
                        lang["content_html"] = inner
                    elif textarea_idx == 1:
                        lang["metadesc_html"] = inner
                    elif textarea_idx == 2:
                        lang["keywords_html"] = inner
                    textarea_idx += 1
 
                i += 1
 
            languages.append(lang)
        else:
            i += 1
 
    return languages
 
 
# ---------------------------------------------------------------------------
# Build a single .docx from one historyElement
# ---------------------------------------------------------------------------
 
def build_docx(languages: list[dict]):
    doc = DocumentFactory()
 
    for idx, lang in enumerate(languages):
        # ── Heading 1: language title ────────────────────────────────────────
        doc.add_heading(lang["title"], level=1)
 
        # ── Content ──────────────────────────────────────────────────────────
        add_html_to_doc(doc, lang["content_html"])
 
        # ── Heading 4: Meta Description ──────────────────────────────────────
        doc.add_heading("Meta Description:", level=4)
        add_html_to_doc(doc, lang["metadesc_html"])
 
        # ── Heading 4: Meta Keywords ─────────────────────────────────────────
        doc.add_heading("Meta Keywords:", level=4)
        add_html_to_doc(doc, lang["keywords_html"])
 
        # ── 3 empty lines between languages (not after the last one) ─────────
        if idx < len(languages) - 1:
            for _ in range(3):
                doc.add_paragraph()
 
    return doc
 
 
# ---------------------------------------------------------------------------
# Filename sanitisation
# ---------------------------------------------------------------------------
 
def sanitize(name: str) -> str:
    """Strip characters that are illegal in Windows/macOS filenames."""
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", name)
    return name.strip()
 
 
# ---------------------------------------------------------------------------
# Main conversion entry point
# ---------------------------------------------------------------------------
 
def convert_file(html_path: str, output_dir: str):
    """Parse *html_path* and write one .docx per historyElement into *output_dir*."""
    with open(html_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")
 
    elements = soup.find_all("div", class_="historyElement")
    if not elements:
        print(f"  [warn] No historyElement divs found in {html_path}")
        return
 
    os.makedirs(output_dir, exist_ok=True)
 
    for rank, element in enumerate(elements, start=1):
        languages = parse_history_element(element)
        if not languages:
            print(f"  [warn] historyElement #{rank}: no language sections found, skipping")
            continue
 
        first_title = languages[0]["title"]
        filename = sanitize(f"{rank}. Başlık {first_title}") + ".docx"
        out_path = os.path.join(output_dir, filename)
 
        doc = build_docx(languages)
        doc.save(out_path)
        print(f"{filename} Completed")
 
 

def main():
    if len(sys.argv) < 2:
        print("Usage: python script.py filename.html [output_dir]")
        sys.exit(1)

    script_dir = Path(__file__).resolve().parent

    input_name = sys.argv[1]
    input_path = script_dir / input_name

    output_dir = sys.argv[2] if len(sys.argv) > 2 else "./output"

    if not input_path.exists() or not input_path.is_file():
        print(f"Error: '{input_name}' not found in script directory.")
        sys.exit(1)

    print(f"Processing {input_path.name} …")
    convert_file(str(input_path), output_dir)

    print("\nDone.")
 
if __name__ == "__main__":
    main()